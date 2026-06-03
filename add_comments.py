"""
add_comments.py - 为 DOCX 文档批量添加批注

基于 docx-comments 库，支持精准文本锚定批注。

支持两种批注目标：
    1. 段落（默认）：通过 match_text 在全文段落中搜索匹配
    2. 单元格（target_type="cell"）：通过 table_index 定位表格，在表格内搜索匹配

批注 JSON 格式：

段落模式（默认）：
[
  {
    "match_text": "工作成果",     # 全文匹配的文本片段
    "match_occurrence": 1,       # 可选，全篇第几次出现，默认 1
    "title": "格式问题",          # 批注标题
    "content": "建议使用列表展示"  # 批注正文
  }
]

单元格模式（通过 table_index 定位表格，在表格内搜索匹配；不指定则遍历所有表格）：
[
  {
    "match_text": "协同消息反馈",
    "target_type": "cell",
    "table_index": 0,            # 可选，表格索引（从 0 开始），不指定则遍历所有表格
    "match_occurrence": 1,       # 可选，第几次出现，默认 1
    "title": "内容检查",
    "content": "请确认此功能描述是否完整。"
  }
]

使用示例：
  python add_comments.py input.docx output.docx annotations.json --author 张三
  python add_comments.py input.docx --list-paragraphs
  python add_comments.py input.docx output.docx annotations.json --author 张三 --list-tables
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx_comments import CommentManager, PersonInfo


def _get_run_text(run) -> str:
    """获取单个 run 的文本内容"""
    t = run.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
    return (t.text or "") if t is not None else ""


def _get_para_text(paragraph) -> str:
    """从 lxml 段落元素中提取完整文本（拼接所有 run）"""
    return "".join(_get_run_text(run) for run in paragraph.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
    ))


def _get_para_style(paragraph) -> str:
    """获取段落样式名"""
    pPr = paragraph.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is None:
        return ""
    pStyle = pPr.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
    return pStyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "") if pStyle is not None else ""


def _find_text_runs(paragraph, match_text: str):
    """
    在段落内定位 match_text 的位置，
    返回 (start_run_idx, end_run_idx) 用于锚定批注。

    通过累加各 run 的字符位置来计算。
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    runs = paragraph.findall(f".//{W}r")
    run_texts = []
    for run in runs:
        text = _get_run_text(run)
        run_texts.append(text)

    full_text = "".join(run_texts)

    if not match_text:
        # 空匹配文本 -> 使用第一个非空 run
        for i, rt in enumerate(run_texts):
            if rt.strip():
                return (i, i)
        return (0, 0) if runs else (0, 0)

    # 查找 match_text 在段落中的位置
    pos = full_text.find(match_text)
    if pos == -1:
        return None, None

    match_start = pos
    match_end = pos + len(match_text)

    # 计算 match_start 和 match_end 落在哪个 run 中
    char_offset = 0
    start_run_idx = 0
    end_run_idx = 0
    found_start = False

    for i, rt in enumerate(run_texts):
        run_start = char_offset
        run_end = char_offset + len(rt)
        if rt:
            if not found_start and run_start <= match_start < run_end:
                start_run_idx = i
                end_run_idx = i
                found_start = True
            elif found_start and run_start <= match_end <= run_end:
                end_run_idx = i
                break
            elif found_start and match_end > run_end:
                end_run_idx = i
        char_offset = run_end

    return start_run_idx, end_run_idx


def _get_cell_text(cell) -> str:
    """获取单元格内所有段落拼接的文本"""
    return "\n".join(p.text for p in cell.paragraphs)


def _find_text_runs_in_cell(cell, match_text: str, occurrence: int = 1):
    """
    在单元格内定位 match_text 的位置，
    返回 (anchor_para, start_run_idx, end_run_idx) 用于锚定批注。

    在单元格所有段落拼接的文本中搜索第 occurrence 次出现的 match_text，
    然后定位它落在哪个段落的哪个 run 范围内。
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    runs_by_para = []
    para_texts = []

    for para in cell.paragraphs:
        para_elem = para._element
        runs = para_elem.findall(f".//{W}r")
        run_texts = []
        for run in runs:
            text = _get_run_text(run)
            run_texts.append(text)
        runs_by_para.append((runs, run_texts))
        para_texts.append("".join(run_texts))

    # 拼接所有段落文本（用 \n 分隔）
    full_text = "\n".join(para_texts)

    if not match_text:
        # 空匹配文本 -> 使用第一个非空 run
        for runs, run_texts in runs_by_para:
            for i, rt in enumerate(run_texts):
                if rt.strip():
                    return para_to_anchor_para(cell, runs), i, i
            # 段落用 \n 分隔，\n 也算匹配
            if runs and len(run_texts) == 0:
                return para_to_anchor_para(cell, runs), 0, 0
        return (cell.paragraphs[0] if cell.paragraphs else None, 0, 0)

    # 在完整文本中查找第 occurrence 次出现
    pos = 0
    found_count = 0
    match_start = -1
    while pos <= len(full_text):
        idx = full_text.find(match_text, pos)
        if idx == -1:
            break
        found_count += 1
        if found_count == occurrence:
            match_start = idx
            break
        pos = idx + 1

    if match_start == -1:
        return None, None, None

    match_end = match_start + len(match_text)

    # 确定 match_start 落在哪个段落中
    # 段落间用 \n 分隔
    para_offsets = []
    char_offset = 0
    for i, pt in enumerate(para_texts):
        para_offsets.append((i, char_offset, char_offset + len(pt)))
        char_offset += len(pt) + 1  # +1 for \n

    para_idx = None
    para_local_start = 0
    para_local_end = 0

    for pidx, offset, end in para_offsets:
        if offset <= match_start < end:
            para_idx = pidx
            para_local_start = match_start - offset
            para_local_end = match_end - offset
            break

    if para_idx is None:
        return None, None, None

    runs, run_texts = runs_by_para[para_idx]
    if not runs:
        return None, None, None

    anchor_para = para_to_anchor_para(cell, runs)
    if anchor_para is None:
        return None, None, None

    # 在段落的 run 中定位
    char_offset = 0
    start_run_idx = 0
    end_run_idx = 0
    found_start = False

    for i, rt in enumerate(run_texts):
        run_start = char_offset
        run_end = char_offset + len(rt)
        if rt:
            if not found_start and run_start <= para_local_start < run_end:
                start_run_idx = i
                end_run_idx = i
                found_start = True
            elif found_start and run_start <= para_local_end <= run_end:
                end_run_idx = i
                break
            elif found_start and para_local_end > run_end:
                end_run_idx = i
        char_offset = run_end

    return anchor_para, start_run_idx, end_run_idx


def para_to_anchor_para(cell, runs):
    """从 cell 的 runs 找到对应的 python-docx Paragraph 对象"""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for para in cell.paragraphs:
        para_elem = para._element
        para_runs = para_elem.findall(f".//{W}r")
        if para_runs and para_runs[0] is runs[0]:
            return para
    # fallback: 返回第一个段落
    return cell.paragraphs[0] if cell.paragraphs else None


def list_paragraphs(docx_path: str):
    """列出文档所有段落及其样式"""
    doc = Document(docx_path)

    STYLE_LABEL = {
        "Title": "标题",
        "Heading1": "标题1",
        "Heading2": "标题2",
        "Heading3": "标题3",
        "Heading4": "标题4",
        "Heading5": "标题5",
        "Heading6": "标题6",
        "Heading7": "标题7",
        "Heading8": "标题8",
        "Heading9": "标题9",
        "Normal": "正文",
    }

    print(f"\n📄 段落列表（共 {len(doc.paragraphs)} 段）：")
    print(f"{'索引':>5}  {'样式':<12}  {'内容（前60字）'}")
    print("─" * 80)

    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ""
        label = STYLE_LABEL.get(style, style if style else "正文")
        # text = para.text[:60] + "..." if len(para.text) > 60 else para.text
        text = para.text
        print(f"  {i:>4}  {label:<12}  {text}")

    print(f"\n💡 使用 --list-paragraphs 查看段落，配合 --author 参数添加批注")


def list_tables(docx_path: str):
    """列出文档所有表格及其内容"""
    doc = Document(docx_path)

    print(f"\n📄 表格列表（共 {len(doc.tables)} 个）：")
    print(f"{'索引':>5}  {'行数':>5}  {'列数':>5}")
    print("─" * 80)

    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        lines = []
        for j, row in enumerate(table.rows):
            # if j >= 3:
            #     break
            # cells_text = " | ".join(cell.text[:20] for cell in row.cells if cell.text)
            cells_text = " | ".join(cell.text for cell in row.cells)
            lines.append(cells_text)
        text = "\n".join(lines)
        # if len(text) > 120:
        #     text = text[:120] + "..."
        print(f"\n\nTABLE_INDEX: {i:>4} ROW_INDEX: {rows:>4} COL_INDEX: {cols:>4} \n{text}")

    print(f"\n💡 使用 --list-tables 查看表格，配合 target_type=\"cell\" 在 JSON 中添加单元格批注")


def add_comments_from_json(
    input_docx: str,
    output_docx: str,
    annotations: list,
    author: str = "Reviewer",
):
    """
    基于 docx-comments 库，为 DOCX 文档添加批注。

    支持两种批注目标：
    1. 段落（默认）：通过 match_text 在段落中搜索
2. 单元格（target_type="cell"）：通过 table_index 定位表格，在表格内搜索匹配文本

    参数
    ----
    input_docx       : 输入 docx 路径
    output_docx      : 输出 docx 路径
    annotations      : 批注列表
    author           : 批注作者
    """
    input_path = Path(input_docx)
    if not input_path.exists():
        raise FileNotFoundError(f"找不到文件: {input_docx}")

    print(f"加载文档: {input_path.name}")
    doc = Document(str(input_path))

    mgr = CommentManager(doc)

    # 确保作者存在于 people.xml 中
    try:
        person = mgr.get_person(author)
    except KeyError:
        person = mgr.ensure_person(author)

    initials = author[0].upper() if author else "R"

    total = len(annotations)
    success_count = 0

    print(f"准备插入 {total} 条批注 ...")

    for cid, ann in enumerate(annotations):
        match_text = ann.get("match_text", "")
        occurrence = ann.get("match_occurrence", 1)
        title = ann.get("title", f"Comment {cid}")
        content = ann.get("content", "")
        target_type = ann.get("target_type", "paragraph")

        # 校验必填字段
        if not match_text:
            print(f"  [SKIP] comment {cid}: 缺少 match_text")
            continue

        # 构建批注内容
        comment_text = f"{title}: {content}" if title and content else (title or content)

        try:
            if target_type == "cell":
                # 单元格批注模式：通过 table_index 定位表格，在表格内搜索匹配；未指定则遍历所有表格
                table_index = ann.get("table_index")
                if table_index is not None:
                    if table_index < 0 or table_index >= len(doc.tables):
                        print(f"  [SKIP] comment {cid}: table_index={table_index} 超出范围（文档共 {len(doc.tables)} 个表格）")
                        continue
                    tables_to_search = [table_index]
                else:
                    tables_to_search = range(len(doc.tables))

                matched_count = 0
                for t_idx in tables_to_search:
                    table = doc.tables[t_idx]
                    for r_idx, row in enumerate(table.rows):
                        for c_idx, cell in enumerate(row.cells):
                            cell_text = _get_cell_text(cell)
                            if match_text in cell_text:
                                anchor_para, start_run, end_run = _find_text_runs_in_cell(cell, match_text, occurrence)
                                if anchor_para is None:
                                    continue
                                comment_id = mgr.add_comment(
                                    paragraph=anchor_para,
                                    start_run=start_run,
                                    end_run=end_run,
                                    text=comment_text,
                                    author=person,
                                    initials=initials,
                                )
                                matched_count += 1
                                print(f"  [✓] comment {cid}: table[{t_idx}], row[{r_idx}], col[{c_idx}] '{match_text}' → {author}")
                success_count += matched_count
                if matched_count == 0:
                    if table_index is not None:
                        print(f"  [SKIP] comment {cid}: 在 table[{table_index}] 中找不到 '{match_text}'")
                    else:
                        print(f"  [SKIP] comment {cid}: 所有表格中找不到 '{match_text}'")

            else:
                # 段落批注模式（原有逻辑）
                para_idx = None
                target_para = None
                para_text = None
                found_count = 0

                for idx, para in enumerate(doc.paragraphs):
                    text = para.text
                    if match_text in text:
                        found_count += 1
                        if found_count == occurrence:
                            para_idx = idx
                            target_para = para
                            para_text = text
                            break

                if target_para is None:
                    print(f"  [SKIP] comment {cid}: 全文中找不到第 {occurrence} 次出现的 '{match_text}'")
                    continue

                # 对于 lxml 段落元素，需要用 XML 方式获取 runs
                # python-docx Paragraph 的 .element 属性是 lxml Element
                para_elem = target_para._element if hasattr(target_para, '_element') else target_para.element

                # 计算 run 索引范围用于锚定
                start_run, end_run = _find_text_runs(para_elem, match_text)

                if start_run is None:
                    print(f"  [SKIP] comment {cid}: 无法定位文本 '{match_text}' 的 run 位置")
                    continue

                comment_id = mgr.add_comment(
                    paragraph=target_para,
                    start_run=start_run,
                    end_run=end_run,
                    text=comment_text,
                    author=person,
                    initials=initials,
                )
                print(f"  [✓] comment {cid}: [{para_idx}] '{match_text}' → {author}")
                success_count += 1

        except Exception as e:
            print(f"  [✗] comment {cid}: 添加失败 - {e}")

    # 保存文档
    print(f"\n保存输出: {output_docx}")
    doc.save(output_docx)

    print(f"\n✅ 完成！成功插入 {success_count}/{total} 条批注 → {output_docx}")


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="基于 docx-comments 为 DOCX 文档批量添加批注",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
使用示例：
  # 列出文档段落
  python add_comments.py document.docx --list-paragraphs

  # 列出文档表格
  python add_comments.py document.docx --list-tables

  # 添加段落批注（从 JSON 文件）
  python add_comments.py document.docx output.docx annotations.json --author 张三

  # 添加单元格批注（从 JSON 文件）
  python add_comments.py document.docx output.docx annotations.json --author 张三
   # JSON 中需指定 target_type="cell" 和 table_index，在指定表格内搜索匹配
""",
    )
    parser.add_argument("input_docx", help="输入 docx 文件路径")
    parser.add_argument("output_docx", nargs="?", help="输出 docx 文件路径（--list-paragraphs 或 --list-tables 时可省略）")
    parser.add_argument("annotations_json", nargs="?", help="批注 JSON 文件路径或内联 JSON 字符串")
    parser.add_argument("--author", default="Reviewer", help="批注作者姓名（默认 Reviewer）")
    parser.add_argument("--list-paragraphs", action="store_true", help="列出文档段落结构后退出")
    parser.add_argument("--list-tables", action="store_true", help="列出文档表格结构后退出")

    args = parser.parse_args()

    if args.list_paragraphs:
        list_paragraphs(args.input_docx)
        sys.exit(0)

    if args.list_tables:
        list_tables(args.input_docx)
        sys.exit(0)

    if not args.output_docx or not args.annotations_json:
        parser.error("添加批注时需要提供 output_docx 和 annotations_json")

    # 解析批注数据
    ann_input = args.annotations_json
    try:
        ann_path = Path(ann_input)
        annotations = json.loads(
            ann_path.read_text(encoding="utf-8") if ann_path.exists() else ann_input
        )
    except (json.JSONDecodeError, OSError) as e:
        print(f"❌ 无法解析批注数据: {e}")
        sys.exit(1)

    if not isinstance(annotations, list):
        print("❌ 批注数据必须是 JSON 数组格式")
        sys.exit(1)

    add_comments_from_json(
        input_docx=args.input_docx,
        output_docx=args.output_docx,
        annotations=annotations,
        author=args.author,
    )


if __name__ == "__main__":
    main()
