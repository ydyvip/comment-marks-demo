"""
生成一份用于演示的示例 DOCX 文档
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / "scripts"))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_docx(output_path: str = "sample.docx"):
    doc = Document()

    # 标题
    title = doc.add_heading("示例文档：季度工作报告", level=1)

    # 第一段
    p1 = doc.add_paragraph(
        "本报告总结了本季度的主要工作成果与存在的问题。"
        "团队在产品研发、客户服务和市场推广三个方面取得了显著进展，"
        "但在成本控制和人员培训方面仍有较大改进空间。"
    )

    doc.add_heading("一、工作成果", level=2)

    p2 = doc.add_paragraph(
        "本季度共完成新功能开发 12 项，修复线上 Bug 47 个，"
        "用户满意度评分从上季度的 82 分提升至 89 分。"
        "客户投诉率下降 15%，续约率维持在 94% 的高位。"
    )

    doc.add_heading("二、存在问题", level=2)

    p3 = doc.add_paragraph(
        "项目交付周期偏长，平均延期 3.2 天。"
        "新员工培训体系尚不完善，导致入职后 1 个月内离职率偏高。"
        "跨部门协作流程需要进一步优化，减少沟通成本。"
    )

    doc.add_heading("三、下季度计划", level=2)

    p4 = doc.add_paragraph(
        "计划推出 AI 辅助功能模块，预计可将开发效率提升 20%。"
        "完善员工培训手册，建立导师制度，目标将新员工留存率提升至 85%。"
        "优化跨部门协作工具，引入项目管理平台统一管理任务进度。"
    )

    doc.save(output_path)
    print(f"✅ 示例文档已生成：{output_path}")

    # 打印全文及字符索引，方便配置 annotations JSON
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    print("\n文档全文（含换行符）：")
    print("-" * 60)
    for i, ch in enumerate(full_text[:300]):
        if ch == "\n":
            print(f"  [{i:>3}] \\n")
        else:
            print(f"  [{i:>3}] {ch}", end="")
        if i % 10 == 9:
            print()
    print("\n... (只显示前 300 字符)")


if __name__ == "__main__":
    output = sys.argv[1] if len(sys.argv) > 1 else "sample.docx"
    create_sample_docx(output)
