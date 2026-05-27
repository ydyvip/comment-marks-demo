"""
docx_to_json.py - 抽取DOCX文档的段落和表格单元格内容并导出为JSON文件

功能：
- 读取DOCX文档的所有段落内容
- 读取DOCX文档的所有表格单元格内容
- 导出为JSON格式，包含段落和单元格信息
- 支持段落和单元格样式信息（标题、正文、单元格等）

使用方法：
python docx_to_json.py input.docx output.json
"""

import json
import sys
from pathlib import Path
from docx import Document


def extract_document_content(docx_path: str, output_path: str = None):
    """
    从DOCX文档中提取所有段落和表格单元格内容并保存为JSON
    
    Args:
        docx_path: 输入的DOCX文件路径
        output_path: 输出的JSON文件路径（可选）
    
    Returns:
        dict: 包含文档信息的字典
    """
    input_path = Path(docx_path)
    if not input_path.exists():
        raise FileNotFoundError(f"找不到文件: {docx_path}")
    
    # 读取DOCX文档
    doc = Document(docx_path)
    
    # 样式映射
    style_mapping = {
        'Title': '标题',
        'Heading1': '标题1', '1': '标题1',
        'Heading2': '标题2', '2': '标题2',
        'Heading3': '标题3', '3': '标题3',
        'Heading4': '标题4', '4': '标题4',
        'Heading5': '标题5', '5': '标题5',
        'Heading6': '标题6', '6': '标题6',
        'Normal': '正文',
        'List Paragraph': '列表项',
        'Quote': '引用',
        'Code': '代码',
    }
    
    # 提取段落信息
    paragraphs_data = []
    for i, paragraph in enumerate(doc.paragraphs):
        content_info = {
            "index": i,
            "element_type": "paragraph",
            "text": paragraph.text.strip() if paragraph.text else "",
            "style": style_mapping.get(paragraph.style.name, '其他') if paragraph.style else '其他',
            "style_name": paragraph.style.name if paragraph.style else '',
            "text_length": len(paragraph.text) if paragraph.text else 0
        }
        
        # 添加段落属性
        if paragraph.alignment:
            content_info["alignment"] = str(paragraph.alignment)
        
        paragraphs_data.append(content_info)
    
    # 提取表格单元格信息
    cells_data = []
    table_count = 0
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 为每个单元格生成唯一索引
                cell_index = len(paragraphs_data) + len(cells_data)
                cell_content = cell.text.strip() if cell.text else ""
                
                cell_info = {
                    "index": cell_index,
                    "element_type": "cell",
                    "text": cell_content,
                    "style": '单元格',
                    "table_index": table_idx,
                    "row_index": row_idx,
                    "column_index": col_idx,
                    "text_length": len(cell_content) if cell_content else 0,
                    "table_position": f"表格{table_idx+1}-行{row_idx+1}-列{col_idx+1}"
                }
                
                cells_data.append(cell_info)
        table_count += 1
    
    # 合并所有内容
    all_content = paragraphs_data + cells_data
    
    # 构建完整的文档信息
    document_info = {
        "metadata": {
            "file_name": input_path.name,
            "file_path": str(input_path.absolute()),
            "paragraph_count": len(paragraphs_data),
            "table_count": table_count,
            "cell_count": len(cells_data),
            "total_elements": len(all_content),
            "total_characters": sum(len(p["text"]) for p in all_content),
            "total_words": sum(len(p["text"].split()) for p in all_content if p["text"])
        },
        "content": all_content,
        "style_distribution": _get_style_distribution(all_content),
        "summary": {
            "paragraphs": {
                "total": len(paragraphs_data),
                "non_empty": len([p for p in paragraphs_data if p["text"]]),
                "empty": len([p for p in paragraphs_data if not p["text"]]),
                "titles": len([p for p in paragraphs_data if p["style"].startswith('标题')]),
                "max_length": max([len(p["text"]) for p in paragraphs_data if p["text"]], default=0)
            },
            "tables": {
                "total": table_count,
                "cells": {
                    "total": len(cells_data),
                    "non_empty": len([c for c in cells_data if c["text"]]),
                    "empty": len([c for c in cells_data if not c["text"]])
                }
            },
            "element_types": {
                "paragraphs": len([p for p in all_content if p["element_type"] == "paragraph"]),
                "cells": len([p for p in all_content if p["element_type"] == "cell"])
            }
        }
    }
    
    # 输出到JSON文件
    if output_path:
        output_file = Path(output_path)
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(document_info, f, ensure_ascii=False, indent=2)
        print(f"✅ JSON文件已生成: {output_file}")
    
    return document_info


def _get_style_distribution(paragraphs_data):
    """统计段落样式分布"""
    style_count = {}
    for para in paragraphs_data:
        style = para["style"]
        style_count[style] = style_count.get(style, 0) + 1
    return style_count


def print_document_summary(docx_path: str):
    """打印文档摘要信息"""
    try:
        doc = Document(docx_path)
        print(f"\n📄 文档信息：")
        print(f"文件名：{Path(docx_path).name}")
        print(f"段落总数：{len(doc.paragraphs)}")
        
        # 段落统计
        non_empty_para_count = 0
        total_chars = 0
        total_words = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                non_empty_para_count += 1
                total_chars += len(para.text)
                total_words += len(para.text.split())
        
        print(f"非空段落：{non_empty_para_count}")
        print(f"空段落：{len(doc.paragraphs) - non_empty_para_count}")
        
        # 表格统计
        table_count = len(doc.tables)
        total_cells = sum(len(table.rows) * len(table.columns) for table in doc.tables)
        non_empty_cells = sum(1 for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip())
        
        print(f"表格总数：{table_count}")
        print(f"单元格总数：{total_cells}")
        print(f"非空单元格：{non_empty_cells}")
        print(f"空单元格：{total_cells - non_empty_cells}")
        print(f"总字符数：{total_chars}")
        print(f"总单词数：{total_words}")
        
        # 显示前5个段落
        print(f"\n📝 前5个段落：")
        print("-" * 60)
        for i, para in enumerate(doc.paragraphs[:5]):
            preview = para.text[:80].replace("\n", " ") + "..." if len(para.text) > 80 else para.text
            print(f"{i+1:2d}. {preview}")
        
        # 显示表格信息
        if table_count > 0:
            print(f"\n📊 表格信息：")
            print("-" * 40)
            for i, table in enumerate(doc.tables[:3]):  # 只显示前3个表格
                rows = len(table.rows)
                cols = len(table.columns)
                print(f"表格{i+1}: {rows}行 × {cols}列")
            
            if table_count > 3:
                print(f"... 还有{table_count-3}个表格")
            
            # 显示第一个表格的前几行内容
            print(f"\n📋 示例表格（表格1）：")
            sample_table = doc.tables[0]
            print("-" * 40)
            for row_idx, row in enumerate(sample_table.rows[:3]):  # 只显示前3行
                row_data = []
                for col_idx, cell in enumerate(row.cells[:3]):  # 只显示前3列
                    cell_text = cell.text.strip()[:30] + "..." if len(cell.text) > 30 else cell.text.strip()
                    row_data.append(f"列{col_idx+1}: '{cell_text}'")
                print(f"行{row_idx+1}: {' | '.join(row_data)}")
            
            if len(sample_table.rows) > 3:
                print("... 更多行")
            if len(sample_table.columns) > 3:
                print("... 更多列")
        
        print(f"\n💡 提示：现在支持优化的数据结构，使用 docx_to_json.py 生成 JSON 文件")
        print(f"     文件将包含更简洁的结构，并支持便利函数计算派生属性")
            
    except Exception as e:
        print(f"❌ 读取文档失败: {e}")


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("使用方法：")
        print("  python docx_to_json.py input.docx output.json")
        print("  python docx_to_json.py input.docx --summary")
        print("  python docx_to_json.py input.docx")
        print("\n选项：")
        print("  input.docx  - 输入的DOCX文件路径")
        print("  output.json - 输出的JSON文件路径（可选）")
        print("  --summary   - 只显示文档摘要信息")
        return
    
    docx_path = sys.argv[1]
    
    # 检查文件是否存在
    if not Path(docx_path).exists():
        print(f"❌ 文件不存在: {docx_path}")
        return
    
    # 处理不同命令行参数
    if len(sys.argv) == 3:
        if sys.argv[2] == "--summary":
            print_document_summary(docx_path)
        else:
            output_path = sys.argv[2]
            try:
                result = extract_document_content(docx_path, output_path)
                content_count = len(result['content'])
                paragraphs_count = len([c for c in result['content'] if c['element_type'] == 'paragraph'])
                cells_count = len([c for c in result['content'] if c['element_type'] == 'cell'])
                print(f"\n✅ 成功导出 {content_count} 个内容元素到 {output_path}")
                print(f"   - 段落: {paragraphs_count} 个")
                print(f"   - 单元格: {cells_count} 个")
                if 'metadata' in result and result['metadata'].get('table_count', 0) > 0:
                    print(f"   - 表格: {result['metadata']['table_count']} 个")
                print(f"\n💡 提示：现在可以使用新增的便利函数计算派生属性")
            except Exception as e:
                print(f"❌ 导出失败: {e}")
    elif len(sys.argv) == 2:
        # 没有输出路径，生成默认的JSON文件名
        input_path = Path(docx_path)
        default_output = f"{input_path.stem}_optimized.json"
        try:
            result = extract_document_content(docx_path, default_output)
            content_count = len(result['content'])
            paragraphs_count = len([c for c in result['content'] if c['element_type'] == 'paragraph'])
            cells_count = len([c for c in result['content'] if c['element_type'] == 'cell'])
            print(f"\n✅ 成功导出 {content_count} 个内容元素到 {default_output}")
            print(f"   - 段落: {paragraphs_count} 个")
            print(f"   - 单元格: {cells_count} 个")
            if 'metadata' in result and result['metadata'].get('table_count', 0) > 0:
                print(f"   - 表格: {result['metadata']['table_count']} 个")
            print(f"\n💡 提示：现在可以使用新增的便利函数计算派生属性")
        except Exception as e:
            print(f"❌ 导出失败: {e}")
    else:
        print("❌ 参数错误。使用 --summary 获取帮助信息")


if __name__ == "__main__":
    main()